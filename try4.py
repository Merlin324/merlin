import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract

# --- Config OCR ---
OCR_CONFIG = r"-l heb --psm 1"

st.set_page_config(page_title="📖 SHOQED", layout="wide")

st.title("📖 SHOQED")

# --- CSS pour zone de texte RTL, image scrollable, alignement top ---
st.markdown(
    """
    <style>
    textarea[aria-label=""] {
        direction: rtl;
        text-align: right;
        font-family: 'David', 'Arial', sans-serif;
        font-size: 16px;
        height: 600px;
        vertical-align: top;
        caret-color: black;
    }
    .css-1lcbmhc.e1fqkh3o2 {  /* Colonnes alignées en haut */
        align-items: flex-start;
    }
    .scrollable-image {
        max-height: 600px;
        overflow-y: auto;
        border: 1px solid #ddd;
        padding: 5px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Upload image ou PDF ---
uploaded_file = st.file_uploader(
    "Charger une image ou un PDF (première page)", type=["png", "jpg", "jpeg", "pdf"]
)

col1, col2 = st.columns([1, 1], gap="large")

image = None
ocr_text = ""

# --- Colonne gauche : image/PDF ---
with col1:
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            images = convert_from_bytes(uploaded_file.read(), first_page=1, last_page=1)
            image = images[0]
        else:
            image = Image.open(uploaded_file)
        
        st.image(image, use_container_width=True)

        # --- OCR Tesseract ---
        ocr_text = pytesseract.image_to_string(image, config=OCR_CONFIG)
        st.subheader("📝 Texte OCR extrait de l'image/PDF")
        st.text_area("", value=ocr_text, height=200)
    else:
        st.write("📄 Charger une image ou PDF ici")

# --- Colonne droite : texte RTL ---
with col2:
    texte = st.text_area("", value=ocr_text, height=600)  # prérempli avec OCR

    if st.button("💾 Enregistrer en Word"):
        if texte.strip() == "":
            st.warning("Le texte est vide !")
        else:
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run(texte)
            run.font.name = 'David'
            run.font.size = Pt(14)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Télécharger le fichier Word",
                data=buffer,
                file_name="texte_hebreu.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
